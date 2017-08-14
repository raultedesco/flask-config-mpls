
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

head=document.add_heading('Informe Evolutivo')
head.alignment = WD_ALIGN_PARAGRAPH.CENTER
nya='Raul Tedesco'
edad ='35'
p = document.add_paragraph()
p.add_run('APELLIDO Y NOMBRE:').bold = True
p.add_run(nya+'\n')
p.add_run('Edad:').bold = True
p.add_run(edad+'\n')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)


table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'


document.add_page_break()

document.save('demo.docx')






# There is no requirement to use python-docx. I found another Python library for messing with docx files called "paradocx" altought it seems a bit abandoned it works for what I need.

# python-docx would be preferable as the project seems more healthy so a solution based on it is still desired.

# Anyway, here is the paradocx based solution:

# from paradocx import Document
# from paradocx.headerfooter import HeaderPart

# template = 'template.docx'
# newimg = open('new_file.png', 'r')

# doc = Document.from_file(template)
# header = doc.get_parts_by_class(HeaderPart).next()
# img = header.related('http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')[0]

# img.data = newimg.read()
# newimg.close()

# doc.save('prueba.docx')

